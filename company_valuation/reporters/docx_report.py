"""
DOCX Report Generator - creates professional Word documents.
"""

import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from ..core.models import CompanyProfile, ValuationReport, DataSourceType
from ..analyzers.valuation import ValuationAnalyzer


class DocxReportGenerator:
    """Generates professional DOCX valuation reports."""

    def __init__(self):
        self.analyzer = ValuationAnalyzer()

    def generate(
        self,
        report: ValuationReport,
        output_path: str,
        include_raw_data: bool = False
    ) -> str:
        """
        Generate a DOCX report.

        Args:
            report: The valuation report data
            output_path: Path for the output file
            include_raw_data: Whether to include raw data points

        Returns:
            Path to the generated file
        """
        doc = Document()
        profile = report.company

        # Set up styles
        self._setup_styles(doc)

        # Title page
        self._add_title_page(doc, profile)

        # Executive Summary
        self._add_executive_summary(doc, profile)

        # Company Overview
        self._add_company_overview(doc, profile)

        # Valuation Analysis
        self._add_valuation_analysis(doc, profile)

        # Metrics Details
        self._add_metrics_details(doc, profile)

        # Data Sources
        self._add_data_sources(doc, profile, report)

        # Raw Data (optional)
        if include_raw_data:
            self._add_raw_data(doc, profile)

        # Methodology
        self._add_methodology(doc)

        # Disclaimer
        self._add_disclaimer(doc)

        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        # Save document
        doc.save(output_path)
        return output_path

    def _setup_styles(self, doc: Document) -> None:
        """Set up document styles."""
        styles = doc.styles

        # Modify heading styles
        heading1 = styles["Heading 1"]
        heading1.font.size = Pt(18)
        heading1.font.color.rgb = RGBColor(0, 51, 102)
        heading1.font.bold = True

        heading2 = styles["Heading 2"]
        heading2.font.size = Pt(14)
        heading2.font.color.rgb = RGBColor(0, 51, 102)

        # Normal style
        normal = styles["Normal"]
        normal.font.size = Pt(11)
        normal.font.name = "Calibri"

    def _add_title_page(self, doc: Document, profile: CompanyProfile) -> None:
        """Add title page."""
        # Add some spacing
        for _ in range(5):
            doc.add_paragraph()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("COMPANY VALUATION REPORT")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        # Company name
        company = doc.add_paragraph()
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = company.add_run(profile.name or profile.domain)
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(51, 51, 51)

        doc.add_paragraph()

        # Domain
        domain_p = doc.add_paragraph()
        domain_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = domain_p.add_run(profile.domain)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(102, 102, 102)

        # Spacing
        for _ in range(8):
            doc.add_paragraph()

        # Date
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(12)

        # Page break
        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, profile: CompanyProfile) -> None:
        """Add executive summary section."""
        doc.add_heading("Executive Summary", level=1)

        # Valuation highlight box
        if profile.estimated_valuation:
            valuation_str = self.analyzer.format_valuation(profile.estimated_valuation)
            low_str = self.analyzer.format_valuation(profile.valuation_range[0])
            high_str = self.analyzer.format_valuation(profile.valuation_range[1])

            p = doc.add_paragraph()
            run = p.add_run("Estimated Valuation: ")
            run.font.bold = True
            run = p.add_run(valuation_str)
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.font.bold = True

            p = doc.add_paragraph()
            run = p.add_run(f"Valuation Range: {low_str} - {high_str}")
            run.font.size = Pt(12)

            p = doc.add_paragraph()
            run = p.add_run(f"Confidence Score: {profile.confidence_score * 100:.0f}%")
            run.font.size = Pt(12)

        doc.add_paragraph()

        # Key findings
        doc.add_heading("Key Findings", level=2)

        findings = self._generate_key_findings(profile)
        for finding in findings:
            p = doc.add_paragraph(finding, style="List Bullet")

        doc.add_paragraph()

    def _generate_key_findings(self, profile: CompanyProfile) -> list[str]:
        """Generate key findings from profile data."""
        findings = []

        # Company basics
        if profile.name:
            findings.append(f"Company Name: {profile.name}")

        if profile.industry:
            findings.append(f"Industry: {profile.industry}")

        if profile.employee_count:
            findings.append(f"Estimated Employees: {profile.employee_count}")

        if profile.headquarters:
            findings.append(f"Headquarters: {profile.headquarters}")

        if profile.founded_year:
            age = datetime.now().year - profile.founded_year
            findings.append(f"Founded: {profile.founded_year} ({age} years ago)")

        # Data coverage
        source_counts = {}
        for dp in profile.data_points:
            source_name = dp.source_type.value
            source_counts[source_name] = source_counts.get(source_name, 0) + 1

        findings.append(f"Data points collected: {len(profile.data_points)} from {len(source_counts)} sources")

        return findings

    def _add_company_overview(self, doc: Document, profile: CompanyProfile) -> None:
        """Add company overview section."""
        doc.add_heading("Company Overview", level=1)

        # Create info table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        info_fields = [
            ("Company Name", profile.name or "Not determined"),
            ("Domain", profile.domain),
            ("Industry", profile.industry or "Not determined"),
            ("Headquarters", profile.headquarters or "Not determined"),
            ("Employee Count", profile.employee_count or "Not determined"),
            ("Founded", str(profile.founded_year) if profile.founded_year else "Not determined"),
        ]

        for label, value in info_fields:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = value

        doc.add_paragraph()

        # Description
        description = self._get_data_value(profile, "meta_description")
        if description:
            doc.add_heading("About", level=2)
            doc.add_paragraph(description)

        doc.add_page_break()

    def _add_valuation_analysis(self, doc: Document, profile: CompanyProfile) -> None:
        """Add valuation analysis section."""
        doc.add_heading("Valuation Analysis", level=1)

        if profile.valuation_factors:
            # Valuation factors table
            doc.add_heading("Valuation Factors", level=2)

            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            # Header
            headers = ["Factor", "Score", "Weight", "Weighted Score"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

            # Data rows
            total_weighted = 0
            for factor in profile.valuation_factors:
                row = table.add_row()
                weighted_score = factor.score * factor.weight
                total_weighted += weighted_score

                row.cells[0].text = factor.name
                row.cells[1].text = f"{factor.score:.1f}"
                row.cells[2].text = f"{factor.weight * 100:.0f}%"
                row.cells[3].text = f"{weighted_score:.1f}"

            # Total row
            row = table.add_row()
            row.cells[0].text = "Total"
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[3].text = f"{total_weighted:.1f}"
            row.cells[3].paragraphs[0].runs[0].font.bold = True

            doc.add_paragraph()

        # Valuation methodology used
        doc.add_heading("Valuation Methods", level=2)

        methods_used = []

        # Check which methods were used
        market_cap = self._get_data_value(profile, "market_cap")
        if market_cap and market_cap != "N/A":
            methods_used.append(("Market Capitalization", market_cap, "Direct market value (public company)"))

        revenue = self._get_data_value(profile, "revenue_ttm")
        if revenue and revenue != "N/A":
            methods_used.append(("Revenue Multiple", revenue, "Industry-standard revenue multiples applied"))

        employees = profile.employee_count
        if employees:
            methods_used.append(("Employee-Based", f"{employees} employees", "Valuation based on typical per-employee value"))

        funding = self._get_data_value(profile, "total_funding") or self._get_data_value(profile, "funding_amount")
        if funding:
            methods_used.append(("Funding-Based", funding, "Estimated from reported funding rounds"))

        if methods_used:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            headers = ["Method", "Data Point", "Description"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

            for method, data, desc in methods_used:
                row = table.add_row()
                row.cells[0].text = method
                row.cells[1].text = str(data)
                row.cells[2].text = desc
        else:
            doc.add_paragraph("Insufficient data for quantitative valuation methods.")

        doc.add_page_break()

    def _add_metrics_details(self, doc: Document, profile: CompanyProfile) -> None:
        """Add detailed metrics section."""
        doc.add_heading("Detailed Metrics", level=1)

        # Group metrics by category
        categories = {}
        for metric in profile.metrics:
            if metric.category not in categories:
                categories[metric.category] = []
            categories[metric.category].append(metric)

        for category, metrics in categories.items():
            doc.add_heading(category.replace("_", " ").title(), level=2)

            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            headers = ["Metric", "Value", "Unit", "Description"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

            for metric in metrics:
                row = table.add_row()
                row.cells[0].text = metric.name
                row.cells[1].text = f"{metric.value:.1f}" if isinstance(metric.value, float) else str(metric.value)
                row.cells[2].text = metric.unit
                row.cells[3].text = metric.description

            doc.add_paragraph()

    def _add_data_sources(
        self, doc: Document, profile: CompanyProfile, report: ValuationReport
    ) -> None:
        """Add data sources section."""
        doc.add_heading("Data Sources", level=1)

        # Iteration summary
        doc.add_heading("Collection Summary", level=2)

        if report.iterations:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            headers = ["Iteration", "Sources Used", "Data Points", "Duration"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

            for iteration in report.iterations:
                row = table.add_row()
                row.cells[0].text = str(iteration.iteration_number)
                row.cells[1].text = ", ".join(iteration.sources_used[:3])
                row.cells[2].text = str(iteration.data_points_collected)
                row.cells[3].text = f"{iteration.duration_seconds:.1f}s"

        doc.add_paragraph()

        # Source breakdown
        doc.add_heading("Source Breakdown", level=2)

        source_counts = {}
        for dp in profile.data_points:
            source_name = dp.source_type.value
            source_counts[source_name] = source_counts.get(source_name, 0) + 1

        for source, count in sorted(source_counts.items(), key=lambda x: -x[1]):
            doc.add_paragraph(f"{source.replace('_', ' ').title()}: {count} data points", style="List Bullet")

        doc.add_page_break()

    def _add_raw_data(self, doc: Document, profile: CompanyProfile) -> None:
        """Add raw data appendix."""
        doc.add_heading("Appendix: Raw Data", level=1)

        # Group by source type
        by_source = {}
        for dp in profile.data_points:
            source = dp.source_type.value
            if source not in by_source:
                by_source[source] = []
            by_source[source].append(dp)

        for source, data_points in by_source.items():
            doc.add_heading(source.replace("_", " ").title(), level=2)

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            headers = ["Key", "Value", "Confidence"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

            for dp in data_points[:20]:  # Limit to 20 per source
                row = table.add_row()
                row.cells[0].text = dp.key
                row.cells[1].text = str(dp.value)[:100]  # Truncate long values
                row.cells[2].text = dp.confidence.value

            doc.add_paragraph()

    def _add_methodology(self, doc: Document) -> None:
        """Add methodology section."""
        doc.add_heading("Methodology", level=1)

        methodology_text = """
This valuation report was generated using an automated multi-source data collection and analysis system. The methodology includes:

1. **Data Collection**: Information is gathered iteratively from multiple public sources including company websites, social media profiles, news articles, job postings, and financial databases.

2. **Metric Calculation**: Raw data is processed to calculate standardized metrics across categories including web presence, social media reach, growth indicators, technology sophistication, and financial performance.

3. **Valuation Factors**: Metrics are weighted and combined to produce valuation factors that contribute to the overall company assessment.

4. **Valuation Estimation**: Multiple valuation methods are applied where data permits, including:
   - Market capitalization (for public companies)
   - Revenue multiples (industry-specific)
   - Employee-based estimation
   - Funding-based estimation

5. **Confidence Scoring**: Each data point and the overall valuation are assigned confidence scores based on data source reliability and completeness.
        """

        for paragraph in methodology_text.strip().split("\n\n"):
            doc.add_paragraph(paragraph.strip())

    def _add_disclaimer(self, doc: Document) -> None:
        """Add disclaimer section."""
        doc.add_heading("Disclaimer", level=1)

        disclaimer_text = """
This report is provided for informational purposes only and should not be considered as financial, investment, or legal advice. The valuation estimates contained herein are based on publicly available information and automated analysis, which may be incomplete, outdated, or inaccurate.

Key limitations:
- Data is collected from public sources and may not reflect the company's actual financial position
- Valuation estimates are based on industry averages and may not account for company-specific factors
- Private company valuations are inherently uncertain and can vary significantly
- The analysis does not include due diligence or verification of source data

Users should conduct their own independent research and consult with qualified professionals before making any business or investment decisions based on this report.
        """

        p = doc.add_paragraph(disclaimer_text.strip())
        p.paragraph_format.space_after = Pt(12)

        # Footer
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _get_data_value(self, profile: CompanyProfile, key: str) -> Optional[str]:
        """Get the most recent value for a data point key."""
        for dp in reversed(profile.data_points):
            if dp.key == key:
                return dp.value
        return None
